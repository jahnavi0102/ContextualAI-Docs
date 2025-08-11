# backend/documents/tasks.py

import os
import io
import mimetypes
from django.conf import settings
from .models import Document, DocumentChunk
from django.db.transaction import atomic

# New imports for embeddings and vector database
from pinecone import Pinecone, PodSpec # Updated import for Pinecone client v2.x.x+
from sentence_transformers import SentenceTransformer
import fitz # For PDF parsing (PyMuPDF)
from docx import Document as DocxDocument # For DOCX parsing

# --- Global Initialization for Efficiency ---
# These will be loaded once when the RQ worker starts,
# avoiding redundant loading for each task.

pinecone_index = None
embedding_model = None
pc = None # New Pinecone client instance

try:
    embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
    print("SentenceTransformer model 'all-MiniLM-L6-v2' loaded.")
except Exception as e:
    print(f"Error loading SentenceTransformer model: {e}")
    # Set to None if model loading fails

try:
    # Initialize Pinecone using the new client instantiation
    # Use PodSpec if your index is not serverless, otherwise use ServerlessSpec
    pc = Pinecone(
        api_key=settings.PINECONE_API_KEY,
        environment=settings.PINECONE_ENVIRONMENT
    )
    # Access the index via the Pinecone client instance
    pinecone_index = pc.Index(settings.PINECONE_INDEX_NAME)
    print(f"Pinecone initialized and connected to index: {settings.PINECONE_INDEX_NAME}")
except Exception as e:
    print(f"Error initializing Pinecone: {e}")
    pc = None # Set to None if initialization fails
    pinecone_index = None # Set to None if initialization fails

# --- End Global Initialization ---

def create_chunks_with_overlap(text, chunk_size=1000, overlap=200):
    """Create overlapping chunks to avoid breaking context"""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Avoid breaking words
        if end < len(text) and text[end] != ' ':
            last_space = chunk.rfind(' ')
            if last_space > start + chunk_size // 2:
                chunk = text[start:start + last_space]
                end = start + last_space
        
        if chunk.strip():
            chunks.append(chunk.strip())
        
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks

def process_document_task(document_id):
    """
    Background task to process an uploaded document:
    1. Extract text.
    2. Clean extracted text to remove problematic characters.
    3. Chunk the text with overlap.
    4. Generate vector embeddings.
    5. Store chunks in relational DB and upsert embeddings into Pinecone.
    6. Update document status.
    """
    document = None # Initialize document to None for broader scope
    try:
        document = Document.objects.get(id=document_id)
        print(f"Starting background processing for document: {document.filename}")
        
        # Ensure Pinecone and embedding model are available
        if pinecone_index is None or embedding_model is None:
            error_msg = "Skipping document processing: Pinecone or Embedding model not initialized. Check server logs."
            print(error_msg)
            document.status = 'failed'
            document.metadata['processing_error'] = error_msg
            document.save()
            return

        # Change status to processing
        document.status = 'processing'
        document.save()
        
        # Get the full path to the locally stored file
        file_path = document.file.path
        
        file_content = ""
        file_extension = document.filename.split('.')[-1].lower()

        # Handle different file types for text extraction
        if file_extension == 'pdf':
            doc = fitz.open(file_path)
            for page in doc:
                file_content += page.get_text() or '' # Ensure text is not None
            doc.close() # Close the document after processing
        elif file_extension == 'docx':
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                file_content += paragraph.text + '\n'
        elif file_extension in ['txt', 'md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                file_content = f.read()
        else:
            print(f"Unsupported file type for document {document.id}: {file_extension}")
            document.status = 'failed'
            document.metadata['processing_error'] = f"Unsupported file type: {file_extension}"
            document.save()
            return

        if not file_content.strip(): # Use .strip() to check for meaningful content
            print(f"No text extracted from document {document.id} (or content was empty/whitespace).")
            document.status = 'failed'
            document.metadata['processing_error'] = "No text could be extracted or file was empty."
            document.save()
            return

        # --- NEW: Clean null characters from extracted content ---
        # Replace NUL characters (0x00) which are illegal in PostgreSQL string literals
        file_content = file_content.replace('\x00', '') 
        print(f"Cleaned file content for NUL characters for document: {document.id}")

        # --- Improved Text Chunking with Overlap ---
        chunks = create_chunks_with_overlap(file_content, chunk_size=1000, overlap=200)
        
        # Prepare data for both relational DB (DocumentChunk) and Pinecone
        document_chunks_to_create = []
        vectors_to_upsert = []

        with atomic():
            # Delete existing chunks for this document before creating new ones
            DocumentChunk.objects.filter(document=document).delete()
            print(f"Deleted existing chunks in relational DB for document: {document.id}")

            for i, chunk_content in enumerate(chunks):
                # This check ensures we only process and store non-empty chunks
                if not chunk_content.strip():
                    continue

                # 1. Store chunk in relational database
                document_chunks_to_create.append(
                    DocumentChunk(
                        document=document,
                        content=chunk_content,
                        position=i
                    )
                )

                # 2. Generate Embeddings for the chunk
                embedding = embedding_model.encode([chunk_content])[0].tolist()

                # 3. Prepare vector for Pinecone upsert - Store FULL content
                vector_id = f"doc_{document.id}_chunk_{i}" 
                vectors_to_upsert.append({
                    "id": vector_id,
                    "values": embedding,
                    "metadata": {
                        "document_id": str(document.id),
                        "filename": document.filename,
                        "chunk_position": i,
                        "full_content": chunk_content,  # Store full content for RAG
                        "content_snippet": chunk_content[:500] + "..." if len(chunk_content) > 500 else chunk_content,
                        "source_url": document.file.url,
                        **document.metadata
                    }
                })
            
            # Bulk create DocumentChunk objects in Django's relational database
            if document_chunks_to_create:
                DocumentChunk.objects.bulk_create(document_chunks_to_create)
                print(f"Created {len(document_chunks_to_create)} chunks in relational DB for document {document.id}")

            # 4. Upsert vectors to Pinecone
            if vectors_to_upsert:
                pinecone_index.upsert(vectors=vectors_to_upsert)
                print(f"Upserted {len(vectors_to_upsert)} vectors to Pinecone for document {document.id}")
            else:
                print(f"No valid chunks to upsert to Pinecone for document {document.id}.")

        document.status = 'completed'
        document.save()
        print(f"Finished processing and chunking for document: {document.filename}. {len(chunks)} chunks created and indexed.")

    except Document.DoesNotExist:
        print(f"Document with id {document_id} not found.")
    except Exception as e:
        # Capture and store the error message in document metadata for debugging
        print(f"An error occurred during document processing for document {document_id}: {e}")
        if document:
            document.status = 'failed'
            document.metadata['processing_error'] = str(e) # Store the error
            document.save()


# def process_document_task(document_id):
#     """
#     Background task to process an uploaded document:
#     1. Extract text.
#     2. Clean extracted text to remove problematic characters.
#     3. Chunk the text.
#     4. Generate vector embeddings.
#     5. Store chunks in relational DB and upsert embeddings into Pinecone.
#     6. Update document status.
#     """
#     document = None # Initialize document to None for broader scope
#     try:
#         document = Document.objects.get(id=document_id)
#         print(f"Starting background processing for document: {document.filename}")
        
#         # Ensure Pinecone and embedding model are available
#         if pinecone_index is None or embedding_model is None:
#             error_msg = "Skipping document processing: Pinecone or Embedding model not initialized. Check server logs."
#             print(error_msg)
#             document.status = 'failed'
#             document.metadata['processing_error'] = error_msg
#             document.save()
#             return

#         # Change status to processing
#         document.status = 'processing'
#         document.save()
        
#         # Get the full path to the locally stored file
#         file_path = document.file.path
        
#         file_content = ""
#         file_extension = document.filename.split('.')[-1].lower()

#         # Handle different file types for text extraction
#         if file_extension == 'pdf':
#             doc = fitz.open(file_path)
#             for page in doc:
#                 file_content += page.get_text() or '' # Ensure text is not None
#             doc.close() # Close the document after processing
#         elif file_extension == 'docx':
#             doc = DocxDocument(file_path)
#             for paragraph in doc.paragraphs:
#                 file_content += paragraph.text + '\n'
#         elif file_extension in ['txt', 'md']:
#             with open(file_path, 'r', encoding='utf-8') as f:
#                 file_content = f.read()
#         else:
#             print(f"Unsupported file type for document {document.id}: {file_extension}")
#             document.status = 'failed'
#             document.metadata['processing_error'] = f"Unsupported file type: {file_extension}"
#             document.save()
#             return

#         if not file_content.strip(): # Use .strip() to check for meaningful content
#             print(f"No text extracted from document {document.id} (or content was empty/whitespace).")
#             document.status = 'failed'
#             document.metadata['processing_error'] = "No text could be extracted or file was empty."
#             document.save()
#             return

#         # --- NEW: Clean null characters from extracted content ---
#         # Replace NUL characters (0x00) which are illegal in PostgreSQL string literals
#         file_content = file_content.replace('\x00', '') 
#         print(f"Cleaned file content for NUL characters for document: {document.id}")

#         # --- Text Chunking ---
#         chunk_size = 1000
#         # Filter out any chunks that become empty after cleaning or are just whitespace
#         chunks = [chunk for chunk in 
#                   (file_content[i:i + chunk_size] for i in range(0, len(file_content), chunk_size))
#                   if chunk.strip()]
        
#         # Prepare data for both relational DB (DocumentChunk) and Pinecone
#         document_chunks_to_create = []
#         vectors_to_upsert = []

#         with atomic():
#             # Delete existing chunks for this document before creating new ones
#             DocumentChunk.objects.filter(document=document).delete()
#             print(f"Deleted existing chunks in relational DB for document: {document.id}")

#             for i, chunk_content in enumerate(chunks):
#                 # This check ensures we only process and store non-empty chunks
#                 if not chunk_content.strip():
#                     continue

#                 # 1. Store chunk in relational database
#                 document_chunks_to_create.append(
#                     DocumentChunk(
#                         document=document,
#                         content=chunk_content,
#                         position=i
#                     )
#                 )

#                 # 2. Generate Embeddings for the chunk
#                 embedding = embedding_model.encode([chunk_content])[0].tolist()

#                 # 3. Prepare vector for Pinecone upsert
#                 vector_id = f"doc_{document.id}_chunk_{i}" 
#                 vectors_to_upsert.append({
#                     "id": vector_id,
#                     "values": embedding,
#                     "metadata": {
#                         "document_id": str(document.id),
#                         "filename": document.filename,
#                         "chunk_position": i,
#                         "content_snippet": chunk_content[:200] + "..." if len(chunk_content) > 200 else chunk_content,
#                         "source_url": document.file.url,
#                         **document.metadata
#                     }
#                 })
            
#             # Bulk create DocumentChunk objects in Django's relational database
#             if document_chunks_to_create:
#                 DocumentChunk.objects.bulk_create(document_chunks_to_create)
#                 print(f"Created {len(document_chunks_to_create)} chunks in relational DB for document {document.id}")

#             # 4. Upsert vectors to Pinecone
#             if vectors_to_upsert:
#                 pinecone_index.upsert(vectors=vectors_to_upsert)
#                 print(f"Upserted {len(vectors_to_upsert)} vectors to Pinecone for document {document.id}")
#             else:
#                 print(f"No valid chunks to upsert to Pinecone for document {document.id}.")

#         document.status = 'completed'
#         document.save()
#         print(f"Finished processing and chunking for document: {document.filename}. {len(chunks)} chunks created and indexed.")

#     except Document.DoesNotExist:
#         print(f"Document with id {document_id} not found.")
#     except Exception as e:
#         # Capture and store the error message in document metadata for debugging
#         print(f"An error occurred during document processing for document {document_id}: {e}")
#         if document:
#             document.status = 'failed'
#             document.metadata['processing_error'] = str(e) # Store the error
#             document.save()
